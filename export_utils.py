from io import BytesIO
from datetime import datetime

from docx import Document
from docx.shared import Pt

from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import SimpleDocTemplate, Paragraph


# ============================================================
# PDF EXPORT
# ============================================================

def export_to_pdf(title: str, content: str):

    buffer = BytesIO()

    doc = SimpleDocTemplate(buffer)

    styles = getSampleStyleSheet()

    story = []

    title_style = styles["Heading1"]
    normal_style = styles["BodyText"]

    story.append(Paragraph(title, title_style))

    story.append(
        Paragraph(
            f"Generated on: {datetime.now().strftime('%d %B %Y, %I:%M %p')}",
            normal_style,
        )
    )

    story.append(Paragraph("<br/><br/>", normal_style))

    lines = content.split("\n")

    for line in lines:

        if line.strip():

            story.append(
                Paragraph(
                    line.replace("\n", "<br/>"),
                    normal_style,
                )
            )

    doc.build(story)

    buffer.seek(0)

    return buffer


# ============================================================
# WORD EXPORT
# ============================================================

def export_to_docx(title: str, content: str):

    document = Document()

    document.add_heading(title, level=1)

    date = datetime.now().strftime("%d %B %Y")

    document.add_paragraph(f"Generated on: {date}")

    document.add_paragraph()

    paragraph = document.add_paragraph()

    run = paragraph.add_run(content)

    run.font.size = Pt(11)

    buffer = BytesIO()

    document.save(buffer)

    buffer.seek(0)

    return buffer
